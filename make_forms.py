# make_forms.py
# Tạo file Word chứa 5 biểu mẫu đánh giá ứng dụng hỗ trợ trẻ dyslexia (tiếng Việt)
# Yêu cầu: pip install python-docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime

# ========== Tiện ích ==========
def setup_styles(doc: Document):
    styles = doc.styles

    # Tiêu đề lớn
    if 'Tiêu đề lớn' not in [s.name for s in styles]:
        style = styles.add_style('Tiêu đề lớn', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        style.font.size = Pt(20)
        style.font.bold = True

    # Tiêu đề mục
    if 'Tiêu đề mục' not in [s.name for s in styles]:
        style = styles.add_style('Tiêu đề mục', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        style.font.size = Pt(16)
        style.font.bold = True

    # Nội dung
    if 'Nội dung' not in [s.name for s in styles]:
        style = styles.add_style('Nội dung', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        style.font.size = Pt(12)

def add_title(doc: Document, text: str):
    p = doc.add_paragraph(text, style='Tiêu đề lớn')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_section(doc: Document, text: str):
    p = doc.add_paragraph(text, style='Tiêu đề mục')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_para(doc: Document, text: str):
    doc.add_paragraph(text, style='Nội dung')

def add_empty_line(doc: Document, n=1):
    for _ in range(n):
        doc.add_paragraph('')

def add_likert_table(doc: Document, items: list, caption: str = None, cols=('1','2','3','4','5')):
    """
    Tạo bảng Likert 1–5: cột = STT | Nội dung | 1 | 2 | 3 | 4 | 5
    """
    if caption:
        add_para(doc, f'Ghi chú: {caption}')
    rows = len(items) + 1
    table = doc.add_table(rows=rows, cols=7)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    hdr[0].text = 'STT'
    hdr[1].text = 'Nội dung nhận xét'
    for i, c in enumerate(cols, start=2):
        hdr[i].text = c

    for idx, itm in enumerate(items, start=1):
        row = table.rows[idx].cells
        row[0].text = str(idx)
        row[1].text = itm
        for j in range(2, 7):
            row[j].text = '□'  # hộp chọn trống
    add_empty_line(doc, 1)

def add_face_scale_table(doc: Document, items: list, caption: str = None):
    """
    Thang mặt cười 5 mức cho học sinh: 😞 🙁 😐 🙂 😄
    """
    if caption:
        add_para(doc, f'Ghi chú: {caption}')
    rows = len(items) + 1
    table = doc.add_table(rows=rows, cols=7)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    hdr[0].text = 'STT'
    hdr[1].text = 'Câu hỏi (chọn 1 hình mặt phù hợp)'
    faces = ['😞','🙁','😐','🙂','😄']
    for i, f in enumerate(faces, start=2):
        hdr[i].text = f

    for idx, itm in enumerate(items, start=1):
        row = table.rows[idx].cells
        row[0].text = str(idx)
        row[1].text = itm
        for j in range(2, 7):
            row[j].text = '□'
    add_empty_line(doc, 1)

def add_signature_block(doc: Document, left_label='Người điền', right_label='Người tiếp nhận'):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3)
    table.cell(0, 0).text = f'{left_label} (ký, ghi rõ họ tên):\n\n\n'
    table.cell(0, 1).text = f'{right_label} (ký, ghi rõ họ tên):\n\n\n'
    add_empty_line(doc, 1)

# ========== Dữ liệu biểu mẫu ==========
# Thang mô tả cho Likert
LIKERT_DESC = "Thang 1–5: 1=Rất không đồng ý, 2=Không đồng ý, 3=Phân vân, 4=Đồng ý, 5=Rất đồng ý."

# 1- Phiếu CHUYÊN GIA (A1–A4, B1–B5, C3, D2, D3)
EXPERT_ITEMS = [
    # A1 – Cấu trúc – tường minh
    "Mục tiêu học tập hiển thị rõ ràng ở mỗi hoạt động (PA/Cards/Reading).",
    "Ví dụ – phản ví dụ trong bài tập được thiết kế tường minh, đúng trọng tâm.",
    "Trình tự nội dung (âm vị → ghép âm → thanh điệu → chính tả → đọc → hiểu) mạch lạc, tăng dần độ khó.",
    # A2 – Đặc thù tiếng Việt
    "Các quy tắc c/k/qu; ng/ngh; g/gh; s/x; ch/tr; n/l được thể hiện chính xác và nhất quán.",
    "Hệ 6 thanh điệu được mã hoá màu ổn định, dễ nhận biết.",
    "Không đưa ngoại lệ gây nhiễu ở cấp độ đầu; ví dụ được kiểm soát.",
    # A3 – Đa giác quan
    "Âm mẫu (TTS) rõ ràng; có tuỳ chọn nghe chậm cho thanh hỏi/ngã/nặng.",
    "Tô màu thanh điệu hỗ trợ đáng kể việc phân biệt dấu.",
    "Emoji/biểu tượng minh hoạ giúp neo nghĩa hiệu quả cho trẻ.",
    # A4 – Cá nhân hoá
    "Có thể lọc luyện theo nhóm lỗi (tag) phù hợp mục tiêu cá nhân.",
    "Gợi ý nội dung theo dữ liệu (tag còn yếu) hoạt động hữu ích.",
    "Tùy chọn phương ngữ (khi có) phù hợp thực tế lớp/địa phương.",
    # B1 – Dễ dùng
    "Giao diện nhất quán; dễ làm quen ngay buổi đầu.",
    "Không cần nhiều hướng dẫn kĩ thuật để sử dụng cơ bản.",
    "Các nút/chức năng chính xuất hiện ở vị trí dễ thấy.",
    # B2 – Khả năng đọc
    "Chữ có dấu hiển thị rõ trên mọi màn hình; không bị 'rụng dấu'.",
    "Giãn dòng/chữ hợp lý, dễ theo dõi bằng mắt.",
    "Độ tương phản màu nền/chữ tốt, không gây mỏi mắt.",
    # B3 – Thao tác an toàn
    "Vùng chạm đủ lớn (≥44 px) giúp hạn chế bấm nhầm.",
    "Phản hồi tức thì, tránh thao tác thừa.",
    "Hiệu ứng trực quan vừa phải, không gây xao nhãng.",
    # B4 – Công cụ hỗ trợ đọc
    "Spotlight (soi dòng) hữu ích và dễ bật/tắt.",
    "Pacer (con trỏ nhịp) dễ điều chỉnh tốc độ phù hợp.",
    "Echo (đọc nhắc) hỗ trợ mô hình hoá nhịp đọc tốt.",
    # B5 – Truy cập
    "Chế độ Dễ đọc (font/giãn cách) cải thiện trải nghiệm thực sự.",
    "Tôn trọng 'giảm chuyển động' (có thể tắt rung/nháy).",
    "Voice UI/hướng dẫn bằng giọng nói có ích cho trẻ.",
    # C3 – Đồng thuận & an toàn tâm lý
    "Ngôn ngữ giao diện tích cực, không phán xét; không xếp hạng công khai.",
    "Quy trình đồng thuận/giải thích rõ ràng khi đánh giá với trẻ.",
    # D2 – Âm thanh
    "Nghe chậm giúp phân biệt hỏi/ngã/nặng tốt hơn.",
    "Ghi âm/nghe lại hoạt động ổn định, dễ thao tác.",
    # D3 – Unicode
    "Hiển thị tiếng Việt (ưỡ, quyển, nghiêng...) luôn chính xác."
]

# 2- Phiếu PHỤ HUYNH (A3, A5, B1–B5, C1–C2, E1–E3)
PARENT_ITEMS = [
    # A3 – Đa giác quan
    "Âm mẫu dễ nghe; có nút nghe chậm khi con gặp khó.",
    "Màu sắc thanh điệu giúp con phân biệt dấu tốt hơn.",
    "Hình/emoji minh hoạ làm con dễ hiểu nghĩa của từ.",
    # A5 – Cảm nhận tiến bộ
    "Sau phiên ngắn, con tự tin hơn khi đọc to.",
    "Con nhớ lại được 3–5 tiếng vừa luyện.",
    "Con hiểu cách làm và ít cần trợ giúp dần.",
    # B1 – Dễ dùng
    "Tôi dễ hướng dẫn con sử dụng ngay lần đầu.",
    "Giao diện nhất quán, ít phải tìm kiếm.",
    "Các nút chính (Nghe, Dễ/Vừa/Khó, Bắt đầu/Kết thúc) rõ ràng.",
    # B2 – Khả năng đọc
    "Cỡ chữ lớn; chữ có dấu rõ; giãn dòng hợp lý.",
    "Con không kêu mỏi mắt hay khó nhìn.",
    "Tương phản hiển thị ổn định cả sáng/tối.",
    # B3 – Thao tác
    "Nút đủ to; con ít bấm nhầm.",
    "Phản hồi sau mỗi thao tác rõ ràng; con hiểu ngay.",
    "Hiệu ứng vừa phải; không làm con xao nhãng.",
    # B4 – Hỗ trợ đọc
    "Spotlight (soi dòng) hữu ích.",
    "Pacer (nhịp đọc) dễ điều chỉnh, không quá nhanh.",
    "Echo (đọc nhắc) giúp con bắt chước giọng/nhịp tốt.",
    # B5 – Truy cập
    "Chế độ Dễ đọc giúp con thoải mái hơn.",
    "Có thể tắt rung/hiệu ứng mạnh khi cần.",
    "Nút 'nghe tên nút' (giữ 0,4s) hữu ích.",
    # C1 – Ẩn danh & tối thiểu dữ liệu
    "Ứng dụng không yêu cầu tên thật/thông tin cá nhân bắt buộc.",
    "Ứng dụng không tự gửi ghi âm của con.",
    "Tôi thấy yên tâm về quyền riêng tư mặc định.",
    # C2 – Quyền kiểm soát dữ liệu
    "Tôi dễ xem/xoá dữ liệu trên thiết bị.",
    "Đồng bộ chỉ bật khi tôi chủ động; có giải thích rõ.",
    "Tôi hiểu nơi dữ liệu được lưu/đồng bộ (nếu bật).",
    # E1 – Hữu ích & Dễ dùng
    "Ứng dụng giúp con cải thiện khả năng đọc.",
    "Ứng dụng dễ dùng đối với gia đình.",
    # E2 – Hứng thú & thời lượng
    "Con hứng thú; 10–15 phút/buổi là vừa sức.",
    "Con sẵn sàng quay lại luyện tiếp vào ngày hôm sau.",
    # E3 – Ý định sử dụng
    "Tôi muốn tiếp tục sử dụng ứng dụng cho con.",
    "Tôi sẵn sàng giới thiệu ứng dụng cho phụ huynh khác."
]

# 3- Phiếu HỌC SINH (thang mặt cười 5 mức)
STUDENT_ITEMS = [
    "Con thấy ứng dụng dễ dùng.",
    "Chữ to, dễ nhìn, dễ đọc.",
    "Con thích bấm 'nghe' để nghe từ/câu mẫu.",
    "Màu sắc dấu (thanh) giúp con nhận biết tốt hơn.",
    "Hình/emoji làm con dễ hiểu nghĩa của từ.",
    "Con thấy phần thẻ (Cards) vui và dễ luyện.",
    "Con thấy phần âm vị (PA) dễ hiểu.",
    "Con thấy phần đọc đoạn (Reading) vừa sức.",
    "Con thấy trò chơi 60 giây thú vị.",
    "Con muốn học lại vào ngày mai."
]

# 4- Phiếu QUAN SÁT (điền số/đánh dấu)
OBS_SECTIONS = {
    "Thông tin phiên": [
        "Ngày/giờ:",
        "Thời lượng (phút):",
        "Thiết bị (Android/iOS/Khác):",
        "Chế độ hiển thị (Bình thường/Dễ đọc/HC):"
    ],
    "Thao tác & hỗ trợ": [
        "Số lần bấm nhầm nút:",
        "Số lần cần trợ giúp kỹ thuật:",
        "Số lần con yêu cầu nghe chậm:",
        "Số lần tạm dừng vì mệt:"
    ],
    "Thời gian theo nhiệm vụ": [
        "PA (phút):",
        "Cards (phút):",
        "Reading (phút):",
        "Game (phút, nếu có):"
    ],
    "Ghi chú hành vi/quan sát": [
        "Mức hứng thú (thấp/vừa/cao):",
        "Biểu hiện mệt mỏi (Không/Có – nhẹ/Có – rõ):",
        "Khó khăn nổi bật (mô tả ngắn):",
        "Đề xuất của điều phối viên:"
    ]
}

# 5- Bản THÔNG TIN – ĐỒNG THUẬN (ngắn)
CONSENT_TEXT = """
BẢN THÔNG TIN – ĐỒNG THUẬN THAM GIA ĐÁNH GIÁ ỨNG DỤNG “ĐỌC DỄ”

Mục đích: Thu thập ý kiến chuyên gia, phụ huynh và học sinh về khả dụng, truy cập, an toàn dữ liệu và tính phù hợp sư phạm của ứng dụng hỗ trợ trẻ khó đọc (dyslexia) tiếng Việt.

Nội dung: Người tham gia sẽ trải nghiệm 10–20 phút (PA, Cards, Reading; tùy chọn Game), sau đó điền phiếu nhận xét. Không có rủi ro đáng kể; có thể dừng bất kỳ lúc nào.

Dữ liệu và riêng tư:
- Không thu thập thông tin nhận dạng bắt buộc. Dữ liệu đánh giá được ẩn danh, chỉ dùng cho mục đích giáo dục/nghiên cứu.
- Không tự động gửi bản ghi âm, không thu ảnh/video.
- Người tham gia có quyền xem, yêu cầu xoá dữ liệu trên thiết bị.
- Đồng bộ chỉ diễn ra khi phụ huynh chủ động bật (opt‑in) và cung cấp URL/SECRET riêng.

Sự tự nguyện: Việc tham gia hoàn toàn tự nguyện; có thể rút lui mà không ảnh hưởng quyền lợi.

ĐỒNG THUẬN:
Tôi đã đọc và hiểu nội dung trên; tôi đồng ý cho bản thân/con em tôi tham gia đánh giá ứng dụng.
Họ tên phụ huynh/giám hộ: ..............................................
Họ tên học sinh: ...........................................................
Ngày ...... tháng ...... năm 20......

Chữ ký phụ huynh/giám hộ: .....................      Chữ ký người thu thập: .....................
"""

# ========== Tạo tài liệu ==========
def build_document():
    doc = Document()
    setup_styles(doc)

    # Bìa
    add_title(doc, "BỘ BIỂU MẪU ĐÁNH GIÁ ỨNG DỤNG “ĐỌC DỄ”")
    add_para(doc, "Mục đích: Thu thập ý kiến chuyên gia, phụ huynh và học sinh sau khi trải nghiệm ứng dụng hỗ trợ trẻ dyslexia.")
    add_para(doc, f"Ngày xuất bản: {datetime.now().strftime('%d/%m/%Y')}")
    add_empty_line(doc, 2)

    # Hướng dẫn chung
    add_section(doc, "Hướng dẫn chung")
    add_para(doc, "• Vui lòng hoàn thành các mục theo thang đo tương ứng. Không có đáp án đúng/sai; mọi ý kiến đều quý giá.")
    add_para(doc, "• Thời lượng đề nghị: 10–20 phút/phiên (PA 5–7 phút; Cards 4–5 phút; Reading 5–8 phút; Game 1 phút – tuỳ chọn).")
    add_para(doc, "• Thang Likert 1–5: 1=Rất không đồng ý, 2=Không đồng ý, 3=Phân vân, 4=Đồng ý, 5=Rất đồng ý.")
    doc.add_page_break()

    # 1- Phiếu CHUYÊN GIA
    add_title(doc, "PHIẾU ĐÁNH GIÁ CHUYÊN GIA")
    add_para(doc, "Đối tượng: giáo viên tiểu học/SLP/chuyên gia âm vị học. " + LIKERT_DESC)
    add_empty_line(doc, 1)
    add_para(doc, "Thông tin người đánh giá: Họ tên .................; Đơn vị .................; Kinh nghiệm (năm) ..........")
    add_empty_line(doc, 1)
    add_likert_table(doc, EXPERT_ITEMS, caption="Đánh dấu 1 lựa chọn mỗi dòng.")
    add_para(doc, "Câu mở 1: Điểm mạnh nổi bật của ứng dụng là gì?")
    add_para(doc, "................................................................................................................")
    add_para(doc, "Câu mở 2: Ba cải tiến ưu tiên cần thực hiện?")
    add_para(doc, "1) ............................................................................................................")
    add_para(doc, "2) ............................................................................................................")
    add_para(doc, "3) ............................................................................................................")
    add_empty_line(doc, 1)
    add_signature_block(doc, "Chuyên gia", "Người thu thập")
    doc.add_page_break()

    # 2- Phiếu PHỤ HUYNH
    add_title(doc, "PHIẾU ĐÁNH GIÁ PHỤ HUYNH")
    add_para(doc, "Đối tượng: phụ huynh có con trải nghiệm ứng dụng. " + LIKERT_DESC)
    add_empty_line(doc, 1)
    add_para(doc, "Thông tin: Họ tên phụ huynh .................; Tên học sinh .................; Lớp/tuổi .................")
    add_empty_line(doc, 1)
    add_likert_table(doc, PARENT_ITEMS, caption="Đánh dấu 1 lựa chọn mỗi dòng.")
    add_para(doc, "Câu mở 1: Con thích nhất điều gì?")
    add_para(doc, "................................................................................................................")
    add_para(doc, "Câu mở 2: Điều khiến con khó chịu/khó dùng nhất?")
    add_para(doc, "................................................................................................................")
    add_empty_line(doc, 1)
    add_signature_block(doc, "Phụ huynh/giám hộ", "Người thu thập")
    doc.add_page_break()

    # 3- Phiếu HỌC SINH
    add_title(doc, "PHIẾU ĐÁNH GIÁ HỌC SINH (thang mặt cười)")
    add_para(doc, "Chọn 1 hình mặt thể hiện cảm nhận của con cho mỗi câu. 😞 Rất không thích | 🙁 Không thích | 😐 Bình thường | 🙂 Thích | 😄 Rất thích")
    add_empty_line(doc, 1)
    add_para(doc, "Họ tên: ..................  Lớp/tuổi: ............")
    add_empty_line(doc, 1)
    add_face_scale_table(doc, STUDENT_ITEMS, caption="Đánh dấu 1 ô ở mỗi hàng.")
    add_para(doc, "Con thích nhất điều gì khi dùng ứng dụng?")
    add_para(doc, "................................................................................................................")
    add_para(doc, "Phần nào làm con thấy khó nhất?")
    add_para(doc, "................................................................................................................")
    add_empty_line(doc, 1)
    add_signature_block(doc, "Học sinh (ký/viết tên)", "Người thu thập")
    doc.add_page_break()

    # 4-Phiếu QUAN SÁT
    add_title(doc, "PHIẾU QUAN SÁT (do điều phối viên ghi)")
    add_para(doc, "Ghi nhanh số liệu thao tác/thời gian và nhận xét hành vi trong phiên trải nghiệm.")
    add_empty_line(doc, 1)
    for sec, items in OBS_SECTIONS.items():
        add_section(doc, sec)
        table = doc.add_table(rows=len(items), cols=2)
        table.style = 'Table Grid'
        for i, it in enumerate(items):
            table.cell(i,0).text = it
            table.cell(i,1).text = "................................................"
        add_empty_line(doc, 1)
    add_signature_block(doc, "Điều phối viên", "Giám sát (nếu có)")
    doc.add_page_break()

    # 5- Thông tin – Đồng thuận
    add_title(doc, "BẢN THÔNG TIN – ĐỒNG THUẬN")
    for line in CONSENT_TEXT.strip().split("\n"):
        add_para(doc, line)
    add_empty_line(doc, 1)

    return doc

if __name__ == "__main__":
    doc = build_document()
    out_name = "Bo_BangHoi_DocDe_Dyslexia.docx"
    doc.save(out_name)
    print(f"Đã tạo file Word: {out_name}")